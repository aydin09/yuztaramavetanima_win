import tkinter as tk
import cv2
from PIL import Image, ImageTk
import Process as p
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt
from docx.shared import Mm, Cm, Inches
from docx.shared import Length


class Home:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tarama ve Tanıma')
        self.master.resizable(False, False)

        # Gets the requested values of the height and widht.
        windowWidth = 800
        windowHeight = 400
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüz Tarama ve Tanımaya Hoşgeldiniz", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=20)

        self.train_bt = tk.Button(self.master, text ="Yüz Tarama", height=2, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_train)
        self.train_bt.pack(fill=tk.X,padx=50, pady=40)

        self.recognition_bt = tk.Button(self.master, text ="Yüz Tanıma", height=2, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_recognition)
        self.recognition_bt.pack(fill=tk.X,padx=50, pady=10)

    def redirect_train(self):
        self.master.destroy()
        root = tk.Tk() 
        GUI = Train(root)

    def redirect_recognition(self):
        self.master.destroy()
        root = tk.Tk() 
        GUI = Recognition(root)

class Train:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tarama')
        self.master.resizable(False, False)

        # Gets the requested values of the height and widht.
        windowWidth = 1000
        windowHeight = 560
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüz Taranıyor", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=10)

        self.imageFrame = tk.Frame(self.master, width=600, height=500)
        self.imageFrame.pack(side=tk.LEFT, padx=10, pady=20)

        self.entryFrame = tk.Frame(self.master, bg="white")
        self.entryFrame.pack(pady=50)

        self.label_title = tk.Label(self.entryFrame, justify=tk.LEFT, text="AdNo (İngilizce karakter) giriniz:", bg="white", font = "Times 15")
        self.label_title.pack(fill=tk.X, pady=10)

        self.label = tk.Text(self.entryFrame, height=1, width=30, font = "Times 15")
        self.label.pack(fill=tk.X, pady=10)

        self.result = tk.Label(self.entryFrame, justify=tk.LEFT, bg="white", font = "Times 15")
        self.result.pack(fill=tk.X, pady=10)

        self.train_bt = tk.Button(self.master, text ="Tarama", height=1, width=10, bg="gray", fg="white", font = "Times 15 bold", command = self.train)
        self.train_bt.pack(pady=10)

        self.back_bt = tk.Button(self.master, text ="Geri", height=1, width=10, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_main)
        self.back_bt.pack(pady=10)

        self.video_capture = cv2.VideoCapture(cv2.CAP_DSHOW)

        self.display = tk.Label(self.imageFrame)
        self.display.pack()

        self.isTrain = False
        self.count = 0
        self.faceCascade = p.train_init()
        self.show_frame()

    def train(self):
        name = self.label.get("1.0",'end-1c')
        self.count = 0

        if name == '':
            self.result['text'] = 'AdNo (İngilizce karakter) giriniz...'
            self.result['fg'] = 'red'
            self.isTrain = False

        else:
            self.isTrain = True
            self.result['fg'] = 'black'
            self.label['state'] = 'disabled'
            self.train_bt['state'] = 'disabled'

    def redirect_main(self):
        self.video_capture.release()
        self.master.destroy()
        root = tk.Tk() 
        GUI = Home(root)

    def show_frame(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        coords, _img = p.detect_face(image, self.faceCascade)

        if self.isTrain == True:
            if self.count < 20:
                self.count = p.process_level(_img, coords, self.label.get("1.0",'end-1c'), self.count)
                self.result['text'] = 'İşleniyor...' + str(self.count * 5) + '%'
            else:
                p.train_classifier()
                self.result['fg'] = 'green'
                self.result['text'] = 'Başarılı!'
                self.label['state'] = 'normal'
                self.label.delete(1.0, tk.END)
                self.train_bt['state'] = 'normal'
                self.isTrain = False

        img = Image.fromarray(_img)
        imgtk = ImageTk.PhotoImage(image=img)
        self.display.imgtk = imgtk
        self.display.configure(image=imgtk)
        self.master.after(100, self.show_frame) 

class Recognition:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tanıma')
        self.master.resizable(False, False)

        self.y=[]
        

        # Gets the requested values of the height and widht.
        windowWidth = 1000
        windowHeight = 560
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüzünüz Tanınıyor", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=10)

        self.imageFrame = tk.Frame(self.master, width=800, height=500)
        self.imageFrame.pack(side=tk.LEFT, padx=10, pady=20)

        self.entryFrame = tk.Frame(self.master, bg="white")
        self.entryFrame.pack(pady=50)

        self.label_title = tk.Label(self.entryFrame, justify=tk.LEFT, text="Adınız:", bg="white", font = "Times 15")
        self.label_title.pack(fill=tk.X, pady=10)

        self.result = tk.Label(self.entryFrame, justify=tk.LEFT, text="",bg="white", font = "Times 15")
        self.result.pack(fill=tk.X, pady=10)

        self.back_bt = tk.Button(self.master, text ="Geri", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_main)
        self.back_bt.pack(pady=10)

        self.yoklama_bt = tk.Button(self.master, text ="Yoklama Al", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.yoklama)
        self.yoklama_bt.pack(pady=10)

        self.yaz_bt = tk.Button(self.master, text ="Yazdır", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.yazdir)
        self.yaz_bt.pack(pady=10)

        self.video_capture = cv2.VideoCapture(cv2.CAP_DSHOW)

        self.display = tk.Label(self.imageFrame)
        self.display.pack()

        self.faceCascade, self.clf, self.labels = p.recognize_init()

        self.show_frame()

        self.metin= tk.Text(self.master, width=20)
        self.metin.pack()  


    def yoklama(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        label, _img = p.recognize_face(image, self.faceCascade, self.clf, self.labels)

        self.y.append(label)

        for h in self.y:
            if h =="":
                self.y.remove(h)
        
        print(set(self.y))

        self.metin.delete('1.0', tk.END)

        for k in set(self.y):
            self.metin.insert(tk.INSERT,k+"\n")
        
    def yazdir(self):
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        paragraph = document.add_paragraph("YOKLAMA LİSTESİ")
        paragraph.alignment = 1

        for yaz in set(self.y):
            paragraph = document.add_paragraph(yaz)

        document.save('yoklama.docx')

        os.startfile("yoklama.docx")
           
    def redirect_main(self):
        self.video_capture.release()
        self.master.destroy()
        root = tk.Tk() 
        GUI = Home(root)

    def show_frame(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        label, _img = p.recognize_face(image, self.faceCascade, self.clf, self.labels)

        self.result['text'] = label

        img = Image.fromarray(_img)
        imgtk = ImageTk.PhotoImage(image=img)
        self.display.imgtk = imgtk
        self.display.configure(image=imgtk)
        self.master.after(100, self.show_frame)

def main():
    root = tk.Tk()
    GUI = Home(root)
    root.mainloop()

if __name__ == '__main__':
    main()
