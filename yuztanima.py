import tkinter as tk
import cv2
import numpy as np
from PIL import Image, ImageTk
import os
import time
import pickle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm

class Home:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tarama ve Tanıma')
        self.master.resizable(False, False)

        # Gets the requested values of the height and widht.
        windowWidth = 800
        windowHeight = 400
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüz Tarama ve Tanımaya Hoşgeldiniz", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=20)

        self.train_bt = tk.Button(self.master, text ="Yüz Tarama", height=2, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_train)
        self.train_bt.pack(fill=tk.X,padx=50, pady=40)

        self.recognition_bt = tk.Button(self.master, text ="Yüz Tanıma", height=2, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_recognition)
        self.recognition_bt.pack(fill=tk.X,padx=50, pady=10)

    def redirect_train(self):
        self.master.destroy()
        root = tk.Tk() 
        GUI = Train(root)

    def redirect_recognition(self):
        self.master.destroy()
        root = tk.Tk() 
        GUI = Recognition(root)

class Train:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tarama')
        self.master.resizable(False, False)

        # Gets the requested values of the height and widht.
        windowWidth = 1000
        windowHeight = 560
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüz Taranıyor", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=10)

        self.imageFrame = tk.Frame(self.master, width=600, height=500)
        self.imageFrame.pack(side=tk.LEFT, padx=10, pady=20)

        self.entryFrame = tk.Frame(self.master, bg="white")
        self.entryFrame.pack(pady=50)

        self.label_title = tk.Label(self.entryFrame, justify=tk.LEFT, text="AdNo (İngilizce karakter) giriniz:", bg="white", font = "Times 15")
        self.label_title.pack(fill=tk.X, pady=10)

        self.label = tk.Text(self.entryFrame, height=1, width=30, font = "Times 15")
        self.label.pack(fill=tk.X, pady=10)

        self.result = tk.Label(self.entryFrame, justify=tk.LEFT, bg="white", font = "Times 15")
        self.result.pack(fill=tk.X, pady=10)

        self.train_bt = tk.Button(self.master, text ="Tarama", height=1, width=10, bg="gray", fg="white", font = "Times 15 bold", command = self.train)
        self.train_bt.pack(pady=10)

        self.back_bt = tk.Button(self.master, text ="Geri", height=1, width=10, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_main)
        self.back_bt.pack(pady=10)

        self.video_capture = cv2.VideoCapture(cv2.CAP_DSHOW)

        self.display = tk.Label(self.imageFrame)
        self.display.pack()

        self.isTrain = False
        self.count = 0
        self.faceCascade = train_init()
        self.show_frame()

    def train(self):
        name = self.label.get("1.0",'end-1c')
        self.count = 0

        if name == '':
            self.result['text'] = 'AdNo (İngilizce karakter) giriniz...'
            self.result['fg'] = 'red'
            self.isTrain = False

        else:
            self.isTrain = True
            self.result['fg'] = 'black'
            self.label['state'] = 'disabled'
            self.train_bt['state'] = 'disabled'

    def redirect_main(self):
        self.video_capture.release()
        self.master.destroy()
        root = tk.Tk() 
        GUI = Home(root)

    def show_frame(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        coords, _img = detect_face(image, self.faceCascade)

        if self.isTrain == True:
            if self.count < 20:
                self.count = process_level(_img, coords, self.label.get("1.0",'end-1c'), self.count)
                self.result['text'] = 'İşleniyor...' + str(self.count * 5) + '%'
            else:
                train_classifier()
                self.result['fg'] = 'green'
                self.result['text'] = 'Başarılı!'
                self.label['state'] = 'normal'
                self.label.delete(1.0, tk.END)
                self.train_bt['state'] = 'normal'
                self.isTrain = False

        img = Image.fromarray(_img)
        imgtk = ImageTk.PhotoImage(image=img)
        self.display.imgtk = imgtk
        self.display.configure(image=imgtk)
        self.master.after(100, self.show_frame) 

class Recognition:
    def __init__(self, master):
        self.master = master
        self.master.title('Yüz Tanıma')
        self.master.resizable(False, False)

        self.y=[]
        
        # Gets the requested values of the height and widht.
        windowWidth = 1000
        windowHeight = 560
 
        # Gets both half the screen width/height and window width/height
        positionRight = int(self.master.winfo_screenwidth()/2 - windowWidth/2)
        positionDown = int(self.master.winfo_screenheight()/2 - windowHeight/2)
 
        # Positions the window in the center of the page.
        self.master.geometry("{}x{}+{}+{}".format(windowWidth, windowHeight, positionRight, positionDown))

        self.master.config(background="#FFFFFF")

        self.title = tk.Label(self.master, justify=tk.CENTER, text="Yüzünüz Tanınıyor", bg="white", font = "Times 30 bold")
        self.title.pack(fill=tk.X, padx=10, pady=10)

        self.imageFrame = tk.Frame(self.master, width=800, height=500)
        self.imageFrame.pack(side=tk.LEFT, padx=10, pady=20)

        self.entryFrame = tk.Frame(self.master, bg="white")
        self.entryFrame.pack(pady=50)

        self.label_title = tk.Label(self.entryFrame, justify=tk.LEFT, text="Adınız:", bg="white", font = "Times 15")
        self.label_title.pack(fill=tk.X, pady=10)

        self.result = tk.Label(self.entryFrame, justify=tk.LEFT, text="",bg="white", font = "Times 15")
        self.result.pack(fill=tk.X, pady=10)

        self.back_bt = tk.Button(self.master, text ="Geri", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.redirect_main)
        self.back_bt.pack(pady=10)

        self.yoklama_bt = tk.Button(self.master, text ="Yoklama Al", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.yoklama)
        self.yoklama_bt.pack(pady=10)

        self.yaz_bt = tk.Button(self.master, text ="Yazdır", height=1, width=15, bg="gray", fg="white", font = "Times 15 bold", command = self.yazdir)
        self.yaz_bt.pack(pady=10)

        self.video_capture = cv2.VideoCapture(cv2.CAP_DSHOW)

        self.display = tk.Label(self.imageFrame)
        self.display.pack()

        self.faceCascade, self.clf, self.labels = recognize_init()

        self.show_frame()

        self.metin= tk.Text(self.master, width=20)
        self.metin.pack()  

    def yoklama(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        label, _img = recognize_face(image, self.faceCascade, self.clf, self.labels)

        self.y.append(label)

        for h in set(self.y):
            if h =="":
                self.y.remove(h)

        print(set(self.y))
        
        self.metin.delete('1.0', tk.END)

        for k in set(self.y):
            self.metin.insert(tk.INSERT,k+"\n")
        
    def yazdir(self):
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        paragraph = document.add_paragraph("YOKLAMA LİSTESİ")
        paragraph.alignment = 1
        
        table = document.add_table(rows=1, cols=2, style = 'Table Grid')

        cell = table.cell(0,0)
        table.cell(0,0).width = Cm(2.0)
        cell.text ="Sıra No"

        cell = table.cell(0,1)
        table.cell(0,1).width = Cm(14.0)
        cell.text = "Adı Soyadı No"

        table = document.add_table(rows=int(len(set(self.y))), cols=2, style = 'Table Grid')
        
        for t in range(0,len(set(self.y))):
            cell = table.cell(t,0)
            table.cell(t,0).width = Cm(2.0)
            cell.text =str(int(t)+1)
            table.cell(t,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell_font = table.cell(t,0).paragraphs[0].runs[0].font
            
            cell = table.cell(t,1)
            table.cell(t,1).width = Cm(14.0)
            cell.text =self.y[t]
            table.cell(t,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell_font = table.cell(t,1).paragraphs[0].runs[0].font
            
        document.save('yoklama.docx')

        os.startfile("yoklama.docx")
           
    def redirect_main(self):
        self.video_capture.release()
        self.master.destroy()
        root = tk.Tk() 
        GUI = Home(root)

    def show_frame(self):
        _, frame = self.video_capture.read()
        frame = cv2.flip(frame, 1)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGBA)
        label, _img = recognize_face(image, self.faceCascade, self.clf, self.labels)

        self.result['text'] = label

        img = Image.fromarray(_img)
        imgtk = ImageTk.PhotoImage(image=img)
        self.display.imgtk = imgtk
        self.display.configure(image=imgtk)
        self.master.after(100, self.show_frame)

def generate_dataset(img, id, img_id):
    path = "data/"+ id

    if not os.path.exists(path):
        os.makedirs(path)

    cv2.imwrite(path + "/" + str(img_id)+".jpg", img)

def train_classifier():
    faces = []
    ids = []
    label_ids = {}
    current_id = 0

    for id in os.listdir("data"):
        path = [os.path.join("data\\" + id, f) for f in os.listdir("data\\" + id)]

        for image in path:
            img = Image.open(image).convert('L')
            image_array = np.array(img, 'uint8')
        
            faces.append(image_array)

            if not id in label_ids:
                label_ids[id] = current_id
                current_id += 1

            ids.append(label_ids[id])

    ids = np.array(ids)

    with open("labels.pickle", "wb") as f:
        pickle.dump(label_ids, f)
    
    clf = cv2.face.LBPHFaceRecognizer_create()
    clf.train(faces, ids)
    clf.save("Classifier.yml")

def draw_rect(img, classifier, scaleFactor, minNeighbors, color, label):
    gray_img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    features = classifier.detectMultiScale(gray_img, scaleFactor, minNeighbors)
    coords = []

    for (x, y, w, h) in features:
        cv2.rectangle(img, (x,y), (x+w, y+h), color, 2)
        cv2.putText(img, label, (x, y-5), cv2.FONT_HERSHEY_SIMPLEX, 0.8, color, 1, cv2.LINE_AA)
        coords = [x, y, w, h]

    return coords

def train_data(img, coords, id, img_id, count):
    if len(coords) == 4:
        face_img = img[coords[1]:coords[1]+coords[3], coords[0]:coords[0]+coords[2]]
        generate_dataset(face_img, id, img_id)
        count = count + 1

    return count

def train_init():
    faceCascade = cv2.CascadeClassifier('haarcascade_frontalface_default.xml')
    return faceCascade

def detect_face(img, faceCascade):
    coords = draw_rect(img, faceCascade, 1.5, 5, (255,255,255), "Yuz")
    return coords, img

def process_level(img, coords, id, count):
    count = train_data(img, coords, id, time.time(), count)
    return count

def recognize_init():
    faceCascade = cv2.CascadeClassifier('haarcascade_frontalface_default.xml')

    clf = cv2.face.LBPHFaceRecognizer_create()
    clf.read('Classifier.yml')

    labels = {}
    with open("labels.pickle", "rb") as f:
        labels_ = pickle.load(f)
        labels = {v:k for k,v in labels_.items()}

    return faceCascade, clf, labels

def recognize_face(img, faceCascade, clf, labels):    
    gray_img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    features = faceCascade.detectMultiScale(gray_img, 1.5, 5)

    label = ''

    for (x, y, w, h) in features:
        cv2.rectangle(img, (x,y), (x+w, y+h), (255, 255, 255), 2)
        id, conf = clf.predict(gray_img[y:y+h, x:x+w])
        if conf >= 45 and conf <= 85:
            label = labels[id]
        else:
            label = ''
        cv2.putText(img, label, (x, y-5), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 1, cv2.LINE_AA)

    return label, img


def main():
    root = tk.Tk()
    GUI = Home(root)
    root.mainloop()

if __name__ == '__main__':
    main()
